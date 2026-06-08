import os
from docx import Document
from pathlib import Path

DATA_DIR = Path(r"c:\Users\DELL\Desktop\Day08_RAG_pipeline_cohort2\data\landing\legal")
DATA_DIR.mkdir(parents=True, exist_ok=True)

laws = {
    "luat-phong-chong-ma-tuy-2021.docx": {
        "title": "LUẬT PHÒNG, CHỐNG MA TÚY (SỐ 73/2021/QH14)",
        "content": """
Điều 1. Phạm vi điều chỉnh
Luật này quy định về phòng, chống ma túy; quản lý người sử dụng trái phép chất ma túy; cai nghiện ma túy; trách nhiệm của cá nhân, gia đình, cơ quan, tổ chức trong phòng, chống ma túy; quản lý nhà nước và hợp tác quốc tế về phòng, chống ma túy.

Điều 2. Giải thích từ ngữ
1. Chất ma túy là chất gây nghiện, chất hướng thần được quy định trong các danh mục chất ma túy do Chính phủ ban hành.
2. Người nghiện ma túy là người sử dụng chất ma túy, thuốc gây nghiện, thuốc hướng thần và bị lệ thuộc vào các chất này.
"""
    },
    "nghi-dinh-105-2021-nd-cp.docx": {
        "title": "NGHỊ ĐỊNH 105/2021/NĐ-CP QUY ĐỊNH CHI TIẾT VÀ HƯỚNG DẪN THI HÀNH MỘT SỐ ĐIỀU CỦA LUẬT PHÒNG, CHỐNG MA TÚY",
        "content": """
Chương I. NHỮNG QUY ĐỊNH CHUNG
Điều 1. Phạm vi điều chỉnh
Nghị định này quy định chi tiết và hướng dẫn thi hành một số điều của Luật Phòng, chống ma túy về: cơ quan chuyên trách phòng, chống tội phạm về ma túy; kiểm soát các hoạt động hợp pháp liên quan đến ma túy và quản lý người sử dụng trái phép chất ma túy.

Điều 2. Đối tượng áp dụng
Nghị định này áp dụng đối với cơ quan, tổ chức, cá nhân có liên quan đến công tác phòng, chống ma túy tại Việt Nam.
"""
    },
    "bo-luat-hinh-su-2015-chuong-xx.docx": {
        "title": "BỘ LUẬT HÌNH SỰ 2015 (SỬA ĐỔI 2017) - CHƯƠNG XX: CÁC TỘI PHẠM VỀ MA TÚY",
        "content": """
Điều 248. Tội sản xuất trái phép chất ma túy
1. Người nào sản xuất trái phép chất ma túy dưới bất kỳ hình thức nào, thì bị phạt tù từ 02 năm đến 07 năm.
2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 07 năm đến 15 năm:
a) Có tổ chức;
b) Phạm tội 02 lần trở lên;
c) Lợi dụng chức vụ, quyền hạn;
d) Lợi dụng danh nghĩa cơ quan, tổ chức.

Điều 249. Tội tàng trữ trái phép chất ma túy
1. Người nào tàng trữ trái phép chất ma túy mà không nhằm mục đích mua bán, vận chuyển, sản xuất trái phép chất ma túy thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 01 năm đến 05 năm:
a) Đã bị xử phạt vi phạm hành chính về hành vi quy định tại Điều này hoặc đã bị kết án về tội này hoặc một trong các tội quy định tại các điều 248, 250, 251 và 252 của Bộ luật này, chưa được xóa án tích mà còn vi phạm.
"""
    }
}

for filename, data in laws.items():
    doc = Document()
    doc.add_heading(data["title"], 0)
    doc.add_paragraph(data["content"].strip())
    
    filepath = DATA_DIR / filename
    doc.save(filepath)
    print(f"Created: {filepath}")
